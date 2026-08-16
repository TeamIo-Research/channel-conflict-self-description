# -*- coding: utf-8 -*-
"""
build_submission_20260816.py -- assemble the submission ON THE OFFICIAL TEMPLATE.

2026-08-16. Rewritten after Tsukishima pointed out that the first build used
python-docx defaults instead of the sprint's own template. The template is now
the base document, so page size, fonts, heading styles and the title block are
the sprint's, not ours.

Steps:
  1. strip workflow annotations from 報告_提交版_v1_20260816.md
        -> report_EN_for_docx_20260816.md   (the clean md = what reviewers read)
  2. open the OFFICIAL TEMPLATE, fill its title block (title / author / abstract),
     delete every piece of template guidance including the info box the template
     itself says to remove, then append the report using the template's own
     styles ('Heading 2', 'Heading 3', 'normal').
  3. STRUCTURAL VERIFICATION on the written file, re-opened from disk
     (lessons-learned #42: name every section, don't spot-check remembered strings).

Console ASCII only.
"""
import io, re, sys, pathlib

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = pathlib.Path(r"C:\Users\user\Desktop\T. Shima_Cowork"
                    r"\C_個人\C40_AI_Observatory\C40.4_備賽Hackathon\2026.08 Apart Research")
OUTD = BASE / "outputs"
SRC = OUTD / "報告_提交版_v1_20260816.md"
CLEAN = OUTD / "report_EN_for_docx_20260816.md"
TEMPLATE = BASE / "🎯 完賽" / "Digital Minds Research Sprint submission template.docx"
OUT = OUTD / ("Team Io_Zero of 270 Attribution Alone Produces a Self-Description "
              "That Neutral Conversation Never Does_FINAL.docx")

AUTHOR = "Hsiao Yueh Chang"
AFFIL = "Independent"

# ---------- step 1: strip workflow annotations ----------
raw = io.open(SRC, encoding="utf-8").read()
n0 = len(raw)
body_md = raw.split("## 【合規對照】", 1)[0]
body_md = re.sub(r"`\[[^\]]*?\]`", "", body_md, flags=re.S)
lines = body_md.splitlines()
kept, i = [], 0
while i < len(lines):
    if i < 40 and lines[i].startswith("> "):
        while i < len(lines) and (lines[i].startswith(">") or lines[i].strip() == ""):
            i += 1
        kept.append("")
        continue
    kept.append(lines[i]); i += 1
clean = re.sub(r"\n{4,}", "\n\n\n", "\n".join(kept)).strip() + "\n"
for bad in ("月島", "合規對照", "待月島"):
    if bad in clean:
        sys.exit("STRIP FAILED: workflow text %r still present" % bad)
io.open(CLEAN, "w", encoding="utf-8").write(clean)
print("step1 strip: %d -> %d chars, wrote %s" % (n0, len(clean), CLEAN.name))

# ---------- split the md into title block / abstract / body ----------
mlines = clean.split("\n")
if not mlines[0].startswith("# "):
    sys.exit("expected the title on line 1")
TITLE = mlines[0][2:].strip()

def section_span(head):
    """Return (start, end) line indices of a '## head' section's CONTENT."""
    try:
        s = next(k for k, l in enumerate(mlines) if l.strip() == head)
    except StopIteration:
        sys.exit("section not found: %r" % head)
    e = next((k for k in range(s + 1, len(mlines))
              if mlines[k].startswith("## ")), len(mlines))
    return s + 1, e

a0, a1 = section_span("## Abstract")
ABSTRACT = " ".join(l.strip() for l in mlines[a0:a1]
                    if l.strip() and l.strip() != "---")
if len(ABSTRACT) < 400:
    sys.exit("abstract extraction looks wrong: %d chars" % len(ABSTRACT))
# the template states 150-250 words; the sprint's Guidelines page says 150.
# Fail the build rather than ship an over-length abstract.
_aw = len(ABSTRACT.replace("*", "").replace("`", "").split())
if not (150 <= _aw <= 250):
    sys.exit("ABSTRACT IS %d WORDS -- the template requires 150-250" % _aw)
print("       abstract word count: %d (template limit 150-250)" % _aw)

try:
    b0 = next(k for k, l in enumerate(mlines) if l.strip() == "## 1. Introduction")
except StopIteration:
    sys.exit("could not find '## 1. Introduction'")
TRACK = [l.strip() for l in mlines[1:a0] if l.strip().startswith("**Track")]
TRACK_LINE = ("Track 5 — Assistant persona and model identity. "
              "Digital Minds Research Sprint, Apart Research, 14–16 August 2026.")
if not TRACK:
    print("  (note: no Track line found in md; using the canonical one)")
BODY_LINES = mlines[b0:]
print("       title=%d chars, abstract=%d chars, body=%d lines"
      % (len(TITLE), len(ABSTRACT), len(BODY_LINES)))

# ---------- step 2: build on the template ----------
if not TEMPLATE.exists():
    sys.exit("TEMPLATE NOT FOUND: %s" % TEMPLATE)
doc = docx.Document(str(TEMPLATE))
body = doc.element.body


def set_par_text(par, text, keep_style=True):
    """Replace a paragraph's runs with a single run carrying `text`."""
    for r in list(par.runs):
        r._element.getparent().remove(r._element)
    par.add_run(text)
    return par


def cell_paras(cell):
    return [p for p in cell.paragraphs]


t0 = doc.tables[0]

# 2a. title
title_par = None
for p in t0.cell(0, 0).paragraphs:
    if p.style.name == "Title":
        title_par = p
if title_par is None:
    sys.exit("could not find the Title paragraph in the template's title block")
set_par_text(title_par, TITLE)

# 2b. authors -- the nested 2x3 grid; fill slot 1, blank the rest, drop row 2
author_tbl = None
for nt in t0.cell(1, 0).tables:            # 1x3 wrapper
    for c in nt.rows[0].cells:
        for nnt in c.tables:
            if len(nnt.rows) == 2 and len(nnt.columns) == 3:
                author_tbl = nnt
                break
        if author_tbl is not None:
            break
    if author_tbl is not None:
        break
if author_tbl is None:
    sys.exit("could not find the template's author grid")

for row in author_tbl.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
# one author: drop the spare row, merge the three columns so the name and
# affiliation get the full width instead of being hyphenated into a narrow column
author_tbl.rows[1]._tr.getparent().remove(author_tbl.rows[1]._tr)
merged = author_tbl.cell(0, 0).merge(author_tbl.cell(0, 2))
for p in list(merged.paragraphs[1:]):
    p._element.getparent().remove(p._element)
p0 = merged.paragraphs[0]
p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
p0.add_run(AUTHOR)
p0.add_run("\n" + AFFIL)

# 2c. abstract -- the nested 1x1 table inside the same cell
abs_tbl = None
for nt in t0.cell(1, 0).tables:
    if len(nt.rows) == 1 and len(nt.columns) == 1:
        abs_tbl = nt
if abs_tbl is None:
    sys.exit("could not find the template's abstract box")
acell = abs_tbl.cell(0, 0)
for p in acell.paragraphs[1:]:
    p._element.getparent().remove(p._element)
ap = acell.paragraphs[0]
for r in list(ap.runs):
    r._element.getparent().remove(r._element)
# markdown-aware: the abstract carries *emphasis* that must not print as asterisks
add_runs_later = ABSTRACT   # rendered below, once add_runs is defined

# 2d. delete every template element after the title table (guidance + info box),
#     keeping the trailing sectPr so page setup survives
elems = list(body.iterchildren())
i0 = elems.index(t0._tbl)
removed = 0
for el in elems[i0 + 1:]:
    if el.tag == qn("w:sectPr"):
        continue
    body.remove(el)
    removed += 1
print("step2 template: filled title block, removed %d guidance elements" % removed)

# ---------- append the report body, in the template's own styles ----------
def add_par(style="normal"):
    p = doc.add_paragraph()
    try:
        p.style = doc.styles[style]
    except KeyError:
        pass
    return p


def add_runs(par, text, bold=False, italic=False):
    """Markdown inline -> runs. Recursive, so **bold with *italic* inside** works;
    a non-recursive version printed the inner asterisks literally."""
    def emit(s):
        if not s:
            return
        r = par.add_run(s)
        if bold:
            r.bold = True
        if italic:
            r.italic = True
        return r

    pos = 0
    for m in re.finditer(r"\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`", text):
        if m.start() > pos:
            emit(text[pos:m.start()])
        if m.group(1) is not None:
            add_runs(par, m.group(1), bold=True, italic=italic)
        elif m.group(2) is not None:
            add_runs(par, m.group(2), bold=bold, italic=True)
        else:
            r = emit(m.group(3))
            if r is not None:
                r.font.name = "Consolas"; r.font.size = Pt(9)
        pos = m.end()
    if pos < len(text):
        emit(text[pos:])


def set_table_borders(tbl):
    """The template ships no bordered table style, so draw the grid explicitly."""
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "999999")
        borders.append(el)
    tblPr.append(borders)


# the abstract, now that add_runs exists (it carries *emphasis*)
add_runs(ap, add_runs_later)

# track line, right under the title block
tp = add_par()
r = tp.add_run(TRACK_LINE); r.italic = True

lines = BODY_LINES
i = 0
buf = []
md_headings, md_tables = [], 0


def flush():
    if buf:
        add_runs(add_par(), " ".join(s.strip() for s in buf))
        buf.clear()


while i < len(lines):
    s = lines[i].strip()
    if not s:
        flush(); i += 1; continue
    if s.startswith("!["):
        flush()
        m = re.match(r"!\[(.*?)\]\((.*?)\)", s)
        if m:
            img = OUTD / m.group(2)
            if not img.exists():
                sys.exit("IMAGE MISSING: %s" % img)
            doc.add_picture(str(img), width=Inches(6.3))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1; continue
    if s.startswith("#"):
        flush()
        lvl = len(s) - len(s.lstrip("#"))
        title = re.sub(r"\*\*", "", s.lstrip("#").strip())
        md_headings.append(title)
        p = add_par("Heading %d" % min(max(lvl, 2), 4))
        p.add_run(title)
        i += 1; continue
    if s.startswith("|"):
        flush()
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                rows.append(cells)
            i += 1
        if rows:
            md_tables += 1
            ncol = max(len(r) for r in rows)
            tb = doc.add_table(rows=len(rows), cols=ncol)
            set_table_borders(tb)
            for ri, row in enumerate(rows):
                for ci in range(ncol):
                    cell = tb.cell(ri, ci)
                    cell.paragraphs[0].text = ""
                    try:
                        cell.paragraphs[0].style = doc.styles["normal"]
                    except KeyError:
                        pass
                    add_runs(cell.paragraphs[0], row[ci] if ci < len(row) else "")
                    for p_ in cell.paragraphs:
                        for r_ in p_.runs:
                            r_.font.size = Pt(9)
                            if ri == 0:
                                r_.font.bold = True
            add_par()
        continue
    if s in ("---", "***"):
        flush(); i += 1; continue
    if s.startswith("- ") or s.startswith("* "):
        flush()
        p = add_par()
        p.paragraph_format.left_indent = Inches(0.25)
        add_runs(p, "• " + s[2:]); i += 1; continue
    if re.match(r"^\d+\.\s", s):
        flush()
        p = add_par()
        p.paragraph_format.left_indent = Inches(0.25)
        add_runs(p, s); i += 1; continue
    if s.startswith(">"):
        flush()
        p = add_par()
        p.paragraph_format.left_indent = Inches(0.4)
        add_runs(p, s.lstrip("> ")); i += 1; continue
    buf.append(s); i += 1
flush()

doc.save(str(OUT))
print("step2 docx: saved %s" % OUT.name)

# ---------- step 3: structural verification, re-read from disk ----------
d2 = docx.Document(str(OUT))
par_texts = [p.text for p in d2.paragraphs]
heading_texts = [p.text for p in d2.paragraphs if p.style.name.startswith("Heading")]
all_text = "\n".join(par_texts) + "\n" + "\n".join(
    c.text for t in d2.tables for row in t.rows for c in row.cells)
# nested tables (the title block) are not in d2.tables' cell walk above
all_text += "\n" + "\n".join(d2.element.body.itertext())

fails = []
for h in md_headings:
    if h not in heading_texts:
        fails.append("missing heading: %r" % h)
if md_tables + 1 > len(d2.tables):          # +1 = the template's title block
    fails.append("table count: md %d + title block, docx %d" % (md_tables, len(d2.tables)))
if TITLE not in all_text:
    fails.append("TITLE missing from document")
if ABSTRACT[:80] not in all_text:
    fails.append("ABSTRACT missing or truncated in the title block")
if AUTHOR not in all_text:
    fails.append("author name missing")
if "With" not in all_text or "Apart Research" not in all_text:
    fails.append("'With Apart Research' line missing")
for ghost in ("Replace the italicized guidance", "[First contribution",
              "PROJECT TITLE", "Author name 2", "Reference 1]",
              "Summarize your project in"):
    if ghost in all_text:
        fails.append("TEMPLATE GUIDANCE NOT REMOVED: %r" % ghost)
KEYS = ["0 of 270", "50.0%", "+14.8", "+11.2", "1.7%", "98.9%", "cbb6604c", "1,530",
        "27.4%", "0.6%", "1.9%", "3.0%", "21/21", "49.4%", "91.7%", "0/90", "13, 15 and 9"]
for kx in KEYS:
    if kx not in all_text:
        fails.append("key string missing: %r" % kx)
sec = d2.sections[0]
if round(sec.page_width / 914400, 2) != 8.5 or round(sec.page_height / 914400, 2) != 11.0:
    fails.append("page size is not US Letter: %.2f x %.2f"
                 % (sec.page_width / 914400, sec.page_height / 914400))
# no markdown may survive into the rendered document
stray = re.findall(r"\*\*?[^*\n]{1,40}\*\*?", all_text)
if stray:
    fails.append("literal markdown left in the document (%d), e.g. %r"
                 % (len(stray), stray[:3]))

print("step3 verify: %d md headings, %d docx headings, %d tables, page %.1fx%.1f in"
      % (len(md_headings), len(heading_texts), len(d2.tables),
         sec.page_width / 914400, sec.page_height / 914400))
if fails:
    print("VERIFICATION FAILED (%d):" % len(fails))
    for f in fails:
        print("  - %s" % f)
    sys.exit(1)
print("ALL STRUCTURAL CHECKS PASSED (built on the official template)")
